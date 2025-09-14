"""File conversion service for processing uploaded documents."""

import shutil
from pathlib import Path
from typing import Optional, Tuple

import magic
import pandas as pd
import pdfplumber
from docx import Document

from .logging import get_logger

logger = get_logger(__name__)


class FileConverter:
    """Service for converting various file types to text format."""

    def __init__(self, uploads_dir: Path, converted_dir: Path):
        """Initialize converter with source and destination directories."""
        self.uploads_dir = Path(uploads_dir)
        self.converted_dir = Path(converted_dir)
        self.converted_dir.mkdir(parents=True, exist_ok=True)

    def detect_file_type(self, file_path: Path) -> str:
        """Detect file type using python-magic."""
        try:
            mime_type = magic.from_file(str(file_path), mime=True)
            return mime_type
        except Exception as e:
            logger.warning(f"Could not detect MIME type for {file_path}: {e}")
            # Fallback to extension-based detection
            extension = file_path.suffix.lower()
            extension_map = {
                ".pdf": "application/pdf",
                ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ".doc": "application/msword",
                ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                ".xls": "application/vnd.ms-excel",
                ".csv": "text/csv",
                ".txt": "text/plain",
                ".json": "application/json",
                ".xml": "application/xml",
            }
            return extension_map.get(extension, "application/octet-stream")

    def process_file(
        self, file_path: Path, original_filename: str
    ) -> Tuple[bool, Optional[str]]:
        """
        Process a file: convert if needed, copy if already text-based.

        Returns:
            Tuple of (success: bool, converted_file_path: Optional[str])
        """
        try:
            mime_type = self.detect_file_type(file_path)
            logger.info(
                f"Processing file {original_filename} with MIME type: {mime_type}"
            )

            # Generate output filename
            base_name = Path(original_filename).stem
            output_path = self.converted_dir / f"{base_name}.txt"

            # Handle different file types
            if mime_type == "application/pdf":
                success = self._convert_pdf(file_path, output_path)
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ]:
                success = self._convert_docx(file_path, output_path)
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
            ]:
                success = self._convert_excel(file_path, output_path)
            elif mime_type == "text/csv":
                success = self._copy_or_convert_csv(file_path, output_path)
            elif mime_type in [
                "text/plain",
                "application/json",
                "application/xml",
                "text/xml",
            ]:
                success = self._copy_text_file(file_path, output_path)
            else:
                logger.warning(
                    f"Unsupported file type: {mime_type} for file {original_filename}"
                )
                return False, None

            if success:
                logger.info(
                    f"Successfully processed {original_filename} -> {output_path.name}"
                )
                return True, str(output_path)
            else:
                logger.error(f"Failed to process {original_filename}")
                return False, None

        except Exception as e:
            logger.error(f"Error processing file {original_filename}: {str(e)}")
            return False, None

    def _convert_pdf(self, input_path: Path, output_path: Path) -> bool:
        """Convert PDF to text using pdfplumber."""
        try:
            with pdfplumber.open(input_path) as pdf:
                text_content = []
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num} ---\n{page_text}\n")

                full_text = "\n".join(text_content)

                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(full_text)

                return True
        except Exception as e:
            logger.error(f"PDF conversion failed: {e}")
            return False

    def _convert_docx(self, input_path: Path, output_path: Path) -> bool:
        """Convert DOCX to text using python-docx."""
        try:
            doc = Document(str(input_path))
            text_content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

            full_text = "\n".join(text_content)

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(full_text)

            return True
        except Exception as e:
            logger.error(f"DOCX conversion failed: {e}")
            return False

    def _convert_excel(self, input_path: Path, output_path: Path) -> bool:
        """Convert Excel to text using openpyxl and pandas."""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(input_path)
            text_content = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(input_path, sheet_name=sheet_name)

                text_content.append(f"--- Sheet: {sheet_name} ---")

                # Convert DataFrame to string representation
                sheet_text = df.to_string(index=False, na_rep="")
                text_content.append(sheet_text)
                text_content.append("")  # Empty line between sheets

            full_text = "\n".join(text_content)

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(full_text)

            return True
        except Exception as e:
            logger.error(f"Excel conversion failed: {e}")
            return False

    def _copy_or_convert_csv(self, input_path: Path, output_path: Path) -> bool:
        """Convert CSV to a more readable text format."""
        try:
            df = pd.read_csv(input_path)
            text_content = df.to_string(index=False)

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text_content)

            return True
        except Exception as e:
            logger.error(f"CSV conversion failed: {e}")
            return False

    def _copy_text_file(self, input_path: Path, output_path: Path) -> bool:
        """Copy text-based files (TXT, JSON, XML) to converted directory."""
        try:
            shutil.copy2(input_path, output_path)
            return True
        except Exception as e:
            logger.error(f"Text file copy failed: {e}")
            return False
